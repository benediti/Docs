# -*- coding: utf-8 -*-
import streamlit as st
import requests
from docx import Document
from docx2pdf import convert
import re
from datetime import datetime
import os
import io
import tempfile
import platform
import shutil
import subprocess
import textwrap

# ========= CONFIGURAÇÃO =========
# Nome do seu arquivo modelo (já editado com tags)
# Usa caminho absoluto baseado no diretório do script
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MODELOS_DISPONIVEIS = {
    "Contrato de Servico": os.path.join(SCRIPT_DIR, "Documento Contrato Serviço - Modelo.docx"),
    "Adendo 2026": os.path.join(SCRIPT_DIR, "Adendo_2026_Modelo.docx"),
}

# API de CNPJ (gratuita)
API_CNPJ = "https://brasilapi.com.br/api/cnpj/v1/"

# ========= FUNÇÕES =========

def numero_para_extenso(valor):
    """
    Retorna valor monetario por extenso com reais e centavos.
    Exemplo: 3400.50 -> 'três mil e quatrocentos reais e cinquenta centavos'
    """
    valor = round(float(valor), 2)

    try:
        from num2words import num2words
        # to='currency' gera reais/centavos corretamente em pt_BR.
        return num2words(valor, lang='pt_BR', to='currency')
    except ImportError:
        inteiro = int(valor)
        centavos = int(round((valor - inteiro) * 100))

        try:
            from num2words import num2words
            parte_inteira = num2words(inteiro, lang='pt_BR')
            parte_centavos = num2words(centavos, lang='pt_BR')
        except Exception:
            parte_inteira = str(inteiro)
            parte_centavos = str(centavos)

        moeda = "real" if inteiro == 1 else "reais"
        if centavos > 0:
            cent = "centavo" if centavos == 1 else "centavos"
            return f"{parte_inteira} {moeda} e {parte_centavos} {cent}"

        return f"{parte_inteira} {moeda}"

def substituir_placeholder_em_paragrafo(paragraph, tag, valor):
    """
    Substitui um placeholder em um parágrafo, lidando com runs quebrados
    """
    # Verificar se o tag está no texto completo
    if tag not in paragraph.text:
        return False
    
    # Juntar todo o texto e verificar posição
    texto_completo = paragraph.text
    
    # Se o placeholder está no texto, vamos reconstruir
    if tag in texto_completo:
        # Limpar todos os runs existentes
        for run in paragraph.runs:
            run.text = ""
        
        # Substituir e adicionar no primeiro run
        novo_texto = texto_completo.replace(tag, str(valor))
        if paragraph.runs:
            paragraph.runs[0].text = novo_texto
        else:
            paragraph.add_run(novo_texto)
        
        return True
    
    return False

def consultar_cnpj(cnpj):
    """
    Consulta dados do CNPJ na API BrasilAPI
    Retorna os dados ou None em caso de erro
    """
    cnpj_limpo = re.sub(r'\D', '', cnpj)
    
    try:
        url = f"{API_CNPJ}{cnpj_limpo}"
        r = requests.get(url, timeout=10)
        
        if r.status_code != 200:
            st.error(f"❌ Erro ao consultar CNPJ (Status {r.status_code})")
            return None
        
        dados = r.json()
        return dados
            
    except requests.exceptions.Timeout:
        st.error("❌ Tempo esgotado ao consultar API. Tente novamente.")
        return None
    except requests.exceptions.RequestException as e:
        st.error(f"❌ Erro na requisição: {str(e)}")
        return None
    except Exception as e:
        st.error(f"❌ Erro inesperado: {str(e)}")
        return None

def converter_docx_para_pdf(docx_bytes, nome_arquivo_base):
    """
    Converte um documento DOCX em bytes para PDF
    Retorna os bytes do PDF e o motor de conversao usado
    """
    def converter_por_convertapi(docx_data):
        """
        Converte DOCX para PDF via ConvertAPI (nuvem), util em ambientes sem Word/LibreOffice.
        Suporta autenticacao por JWT (recomendado) ou Secret (legado).
        """
        def get_secret_value(key, default=""):
            try:
                return st.secrets.get(key, default)
            except Exception:
                return os.getenv(key, default)

        def extrair_jwt(data):
            if not isinstance(data, dict):
                return None

            # Chaves conhecidas
            for key in ("Token", "Jwt", "token", "jwt", "JwtToken", "jwtToken"):
                value = data.get(key)
                if isinstance(value, str) and value.count(".") == 2:
                    return value

            # Fallback: procura qualquer string em formato JWT no payload
            for value in data.values():
                if isinstance(value, str) and value.count(".") == 2:
                    return value

            return None

        def gerar_jwt_convertapi(api_token, kid, expires_in_sec=3600, client_ip=""):
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_token}",
            }
            payload = {
                "Kid": kid,
                "ExpiresInSec": int(expires_in_sec),
            }
            if client_ip:
                payload["ClientIp"] = client_ip

            try:
                resp = requests.post(
                    "https://v2.convertapi.com/token/jwt",
                    json=payload,
                    headers=headers,
                    timeout=30,
                )
                if resp.status_code != 200:
                    st.warning(
                        f"ConvertAPI JWT falhou ({resp.status_code}). Verifique CONVERTAPI_API_TOKEN e CONVERTAPI_KID."
                    )
                    return None

                data = resp.json()
                token = extrair_jwt(data)
                if not token:
                    st.warning("ConvertAPI JWT retornou resposta sem token utilizável.")
                return token
            except Exception as e:
                st.warning(f"Erro ao gerar JWT no ConvertAPI: {str(e)}")
                return None

        secret = ""
        jwt_token = get_secret_value("CONVERTAPI_JWT", "")
        api_token = get_secret_value("CONVERTAPI_API_TOKEN", "") or get_secret_value("CONVERTAPI_TOKEN", "")
        kid = get_secret_value("CONVERTAPI_KID", "") or get_secret_value("CONVERTAPI_JWT_KID", "")
        client_ip = get_secret_value("CONVERTAPI_CLIENT_IP", "")
        expires_in = get_secret_value("CONVERTAPI_EXPIRES_IN_SEC", 3600)

        if not jwt_token and api_token and kid:
            jwt_token = gerar_jwt_convertapi(api_token, kid, expires_in, client_ip)

        try:
            secret = get_secret_value("CONVERTAPI_SECRET", "")
        except Exception:
            secret = ""

        if not jwt_token and not secret:
            st.info("ConvertAPI não configurado (faltam credenciais).")
            return None

        try:
            headers = {}
            params = {}

            if jwt_token:
                headers["Authorization"] = f"Bearer {jwt_token}"
            else:
                params["Secret"] = secret

            response = requests.post(
                "https://v2.convertapi.com/convert/docx/to/pdf",
                params=params,
                headers=headers,
                files={
                    "File": (
                        "documento.docx",
                        docx_data,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                timeout=60,
            )

            if response.status_code != 200:
                detalhe = response.text[:200].replace("\n", " ") if response.text else ""
                st.warning(f"ConvertAPI conversão falhou ({response.status_code}). {detalhe}")
                return None

            payload = response.json()
            arquivos = payload.get("Files", [])
            if not arquivos:
                st.warning("ConvertAPI retornou resposta sem arquivo PDF.")
                return None

            url_pdf = arquivos[0].get("Url")
            if not url_pdf:
                return None

            pdf_response = requests.get(url_pdf, timeout=60)
            if pdf_response.status_code != 200:
                st.warning(f"Download do PDF do ConvertAPI falhou ({pdf_response.status_code}).")
                return None

            return pdf_response.content
        except Exception as e:
            st.warning(f"Erro na integração ConvertAPI: {str(e)}")
            return None

    def gerar_pdf_fallback(docx_data):
        """
        Gera um PDF simples a partir do texto do DOCX quando a conversao nativa falha.
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
        except ImportError:
            st.error("❌ Biblioteca 'reportlab' não encontrada para gerar PDF compatível.")
            st.info("💡 Atualize as dependências com: pip install -r requirements.txt")
            return None

        doc_temp = Document(io.BytesIO(docx_data))
        pdf_buffer = io.BytesIO()
        pdf = canvas.Canvas(pdf_buffer, pagesize=A4)
        largura, altura = A4
        margem = 40
        y = altura - margem

        def escrever_linha(linha, negrito=False):
            nonlocal y
            if y <= margem:
                pdf.showPage()
                y = altura - margem
            fonte = "Helvetica-Bold" if negrito else "Helvetica"
            pdf.setFont(fonte, 10)
            pdf.drawString(margem, y, linha)
            y -= 14

        escrever_linha("Documento gerado automaticamente", negrito=True)
        escrever_linha("", negrito=False)

        for p in doc_temp.paragraphs:
            texto = p.text.strip()
            if not texto:
                continue
            for linha in textwrap.wrap(texto, width=100):
                escrever_linha(linha)

        if doc_temp.tables:
            escrever_linha("", negrito=False)
            escrever_linha("Tabelas", negrito=True)
            for table in doc_temp.tables:
                for row in table.rows:
                    linha_tabela = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
                    if not linha_tabela.strip():
                        continue
                    for linha in textwrap.wrap(linha_tabela, width=100):
                        escrever_linha(linha)

        pdf.save()
        pdf_buffer.seek(0)
        return pdf_buffer.getvalue()

    try:
        # Criar arquivos temporários
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
            tmp_docx.write(docx_bytes)
            tmp_docx_path = tmp_docx.name
        
        tmp_pdf_path = tmp_docx_path.replace('.docx', '.pdf')

        # Conversao principal: Windows + Word (docx2pdf)
        if platform.system() == 'Windows':
            try:
                convert(tmp_docx_path, tmp_pdf_path)
                with open(tmp_pdf_path, 'rb') as f:
                    return f.read(), "Word (docx2pdf)"
            except Exception:
                pass

        # Segunda tentativa: LibreOffice (preserva melhor o layout do Word que o fallback)
        soffice_cmd = shutil.which('soffice')
        if soffice_cmd:
            try:
                subprocess.run(
                    [
                        soffice_cmd,
                        '--headless',
                        '--convert-to',
                        'pdf',
                        '--outdir',
                        os.path.dirname(tmp_docx_path),
                        tmp_docx_path,
                    ],
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                )

                if os.path.exists(tmp_pdf_path):
                    with open(tmp_pdf_path, 'rb') as f:
                        return f.read(), "LibreOffice (soffice)"
            except Exception:
                pass

        # Terceira tentativa: conversao em nuvem (fidelidade melhor que fallback em texto)
        pdf_convertapi = converter_por_convertapi(docx_bytes)
        if pdf_convertapi:
            return pdf_convertapi, "ConvertAPI (nuvem)"

        # Fallback universal para garantir download em PDF
        pdf_fallback = gerar_pdf_fallback(docx_bytes)
        if pdf_fallback:
            return pdf_fallback, "Compativel (fallback)"
        return None, None
            
    except Exception as e:
        st.error(f"❌ Erro ao converter para PDF: {str(e)}")
        return None, None
    finally:
        if 'tmp_docx_path' in locals() and os.path.exists(tmp_docx_path):
            os.unlink(tmp_docx_path)
        if 'tmp_pdf_path' in locals() and os.path.exists(tmp_pdf_path):
            os.unlink(tmp_pdf_path)

def preencher_contrato(tipo_servico, nome_servico, cnpj, ie_cliente, valor, data_inicio,
                       local_execucao, funcoes, observacoes, modelo_path, modelo_nome,
                       dados_cnpj=None):
    """
    Gera um contrato preenchido automaticamente com base em um modelo DOCX.
    Retorna DOCX, PDF e metadados para download.
    """
    
    # --- 1️⃣ Usar dados do CNPJ já consultados ou informados manualmente ---
    if dados_cnpj:
        nome_cliente = dados_cnpj.get("razao_social", "")
        nome_fantasia = dados_cnpj.get("nome_fantasia", "")
        
        # Montar endereço completo
        # A API BrasilAPI não retorna tipo de logradouro separado, apenas o nome
        logradouro = dados_cnpj.get('logradouro', '')  # Ex: BRIGADEIRO FARIA LIMA
        numero = dados_cnpj.get('numero', '')
        complemento = dados_cnpj.get('complemento', '')
        bairro = dados_cnpj.get('bairro', '')
        municipio = dados_cnpj.get('municipio', '')
        uf = dados_cnpj.get('uf', '')
        cep = dados_cnpj.get('cep', '')
        
        # Construir endereço
        partes_endereco = []
        if logradouro:
            partes_endereco.append(logradouro)
        if numero:
            partes_endereco.append(numero)
        if complemento:
            partes_endereco.append(complemento)
        
        endereco_base = ', '.join(partes_endereco)
        
        # Adicionar bairro, cidade e CEP
        if bairro:
            endereco_base += f" - {bairro}"
        if municipio and uf:
            endereco_base += f" - {municipio}/{uf}"
        if cep:
            # Formatar CEP (XXXXX-XXX)
            if len(cep) == 8:
                cep_formatado = f"{cep[:5]}-{cep[5:]}"
                endereco_base += f", CEP {cep_formatado}"
            else:
                endereco_base += f", CEP {cep}"
        
        endereco_cliente = endereco_base
        cnpj_formatado = re.sub(r'\D', '', cnpj)
    else:
        nome_cliente = "NÃO INFORMADO"
        nome_fantasia = ""
        endereco_cliente = "NÃO INFORMADO"
        cnpj_formatado = re.sub(r'\D', '', cnpj)

    # --- 2️⃣ Formatar valor ---
    try:
        # Limpar e converter valor (aceita vírgula ou ponto)
        valor_limpo = str(valor).replace(',', '.').strip()
        valor_float = float(valor_limpo)
        valor_extenso = numero_para_extenso(valor_float)
    except ValueError:
        st.error(f"❌ Valor inválido: '{valor}'. Use formato numérico (ex: 3400.00 ou 3400,00)")
        return None, None, None, None, None

    # --- 3️⃣ Verificar se o modelo existe ---
    if not os.path.exists(modelo_path):
        st.error(f"❌ Erro: Arquivo modelo não encontrado em: {modelo_path}")
        st.info(f"Diretório atual: {os.getcwd()}")
        st.info(f"Arquivos disponíveis: {os.listdir(SCRIPT_DIR) if os.path.exists(SCRIPT_DIR) else 'N/A'}")
        return None, None, None, None, None
    
    # --- 4️⃣ Abrir o modelo e substituir tags ---
    doc = Document(modelo_path)
    
    # Formatar valor monetário
    valor_formatado = f"R$ {valor_float:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    
    substituicoes = {
        "{{tipo_servico}}": tipo_servico,
        "{{nome_servico}}": nome_servico,
        "{{nome_cliente}}": nome_cliente,
        "{{nome_fantasia}}": nome_fantasia if nome_fantasia else "",
        "{{endereco_cliente}}": endereco_cliente,
        "{{cnpj}}": cnpj_formatado,
        "{{ie_cliente}}": ie_cliente,
        "{{funcoes}}": funcoes,
        "{{observacoes}}": observacoes,
        "{{local_execucao}}": local_execucao,
        "{{valor_num}}": valor_formatado,
        "{{valor_extenso}}": valor_extenso.capitalize(),
        "{{data_inicio}}": data_inicio
    }

    # Substituir em parágrafos
    for p in doc.paragraphs:
        for tag, val in substituicoes.items():
            substituir_placeholder_em_paragrafo(p, tag, val)
    
    # Substituir em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for tag, val in substituicoes.items():
                        substituir_placeholder_em_paragrafo(paragraph, tag, val)

    # --- 5️⃣ Salvar DOCX em memória ---
    prefixo_arquivo = "adendo" if "adendo" in modelo_nome.lower() else "contrato"
    nome_base = f"{prefixo_arquivo}_{nome_cliente[:20].strip().replace(' ', '_')}"
    nome_arquivo_docx = f"{nome_base}.docx"
    nome_arquivo_pdf = f"{nome_base}.pdf"
    
    # Salvar DOCX em bytes
    docx_bytes_io = io.BytesIO()
    doc.save(docx_bytes_io)
    docx_bytes_io.seek(0)
    docx_bytes = docx_bytes_io.getvalue()
    
    # --- 6️⃣ Converter para PDF ---
    pdf_bytes, motor_pdf = converter_docx_para_pdf(docx_bytes, nome_base)
    
    return docx_bytes, nome_arquivo_docx, pdf_bytes, nome_arquivo_pdf, motor_pdf

# ========= INTERFACE STREAMLIT =========
def main():
    st.set_page_config(page_title="Gerador de Contratos", page_icon="📄", layout="wide")
    
    st.title("📄 Gerador Automático de Contratos")
    st.markdown("Preencha os dados abaixo para gerar um contrato personalizado automaticamente.")
    
    # Inicializar session state
    if 'dados_cnpj' not in st.session_state:
        st.session_state.dados_cnpj = None
    if 'cnpj_consultado' not in st.session_state:
        st.session_state.cnpj_consultado = ""
    
    # Seção para consulta de CNPJ
    st.markdown("### 🔍 Consulta de CNPJ")
    col_cnpj1, col_cnpj2 = st.columns([3, 1])
    
    with col_cnpj1:
        cnpj_input = st.text_input("Digite o CNPJ para consulta automática", 
                                   value="65035552000180",
                                   help="Digite o CNPJ e clique em Consultar")
    with col_cnpj2:
        st.markdown("<br>", unsafe_allow_html=True)
        btn_consultar = st.button("🔎 Consultar CNPJ", use_container_width=True, type="primary")
    
    # Consultar CNPJ quando o botão for clicado
    if btn_consultar and cnpj_input:
        with st.spinner("Consultando CNPJ..."):
            dados = consultar_cnpj(cnpj_input)
            if dados:
                st.session_state.dados_cnpj = dados
                st.session_state.cnpj_consultado = cnpj_input
                st.success(f"✅ Dados encontrados: {dados.get('razao_social', 'N/A')}")
                
                # Mostrar dados encontrados
                with st.expander("📋 Dados da Empresa", expanded=True):
                    col_info1, col_info2 = st.columns(2)
                    with col_info1:
                        st.write(f"**Razão Social:** {dados.get('razao_social', 'N/A')}")
                        nome_fant = dados.get('nome_fantasia', '')
                        st.write(f"**Nome Fantasia:** {nome_fant if nome_fant else 'Não informado'}")
                        st.write(f"**CNPJ:** {dados.get('cnpj', 'N/A')}")
                    with col_info2:
                        # Montar endereço completo
                        log = dados.get('logradouro', '')
                        num = dados.get('numero', '')
                        comp = dados.get('complemento', '')
                        bairro = dados.get('bairro', '')
                        mun = dados.get('municipio', '')
                        uf = dados.get('uf', '')
                        cep = dados.get('cep', '')
                        
                        # Formatar CEP
                        if cep and len(cep) == 8:
                            cep_formatado = f"{cep[:5]}-{cep[5:]}"
                        else:
                            cep_formatado = cep
                        
                        endereco_partes = []
                        if log:
                            endereco_partes.append(log)
                        if num:
                            endereco_partes.append(num)
                        if comp:
                            endereco_partes.append(comp)
                        
                        endereco_linha1 = ', '.join(endereco_partes)
                        
                        st.write(f"**Endereço:** {endereco_linha1}")
                        st.write(f"**Bairro:** {bairro}")
                        st.write(f"**Cidade/UF:** {mun}/{uf}")
                        st.write(f"**CEP:** {cep_formatado}")
    
    st.markdown("---")
    
    # Formulário principal
    with st.form("contrato_form", clear_on_submit=False):
        st.markdown("### 📝 Dados do Serviço")

        st.caption("Selecione abaixo qual documento deseja gerar.")
        
        col1, col2 = st.columns(2)
        
        with col1:
            tipo_servico = st.text_input("Tipo de Serviço *", 
                                        value="Prestação de Serviços de Limpeza",
                                        help="Ex: Prestação de Serviços de Limpeza")
            
            nome_servico = st.text_input("Nome do Serviço *", 
                                        value="Limpeza e Conservação",
                                        help="Ex: Limpeza e Conservação")
            
            valor = st.text_input("Valor Mensal (R$) *", 
                                 value="3400.00",
                                 help="Valor numérico, ex: 3400.00")
            
            data_inicio = st.text_input("Data de Início *", 
                                       value="03/11/2025",
                                       help="Formato: DD/MM/AAAA")
        
        with col2:
            cnpj = st.text_input("CNPJ *", 
                                value=st.session_state.cnpj_consultado if st.session_state.cnpj_consultado else "65035552000180",
                                help="Será preenchido automaticamente após consulta")
            
            ie_cliente = st.text_input("Inscrição Estadual", 
                                      value="",
                                      help="Deixe em branco se não houver")
            
            local_execucao = st.text_area("Local de Execução *", 
                                         value="Rua Joaquim Murtinho, 225, Bom Retiro - São Paulo/SP",
                                         height=100,
                                         help="Endereço onde o serviço será executado")
        
        st.markdown("### 👥 Funções e Quadro Funcional")
        funcoes = st.text_area("Funções e Quadro Funcional *", 
                              value="Supervisora Operacional / Encarregada – 8h\n4 Auxiliares de Limpeza – 8h",
                              height=120,
                              help="Descreva as funções e quantidade de funcionários")
        
        st.markdown("### 📌 Observações")
        observacoes = st.text_area("Observações", 
                                  value="Serviços de limpeza geral realizados bimestralmente aos sábados.",
                                  height=100,
                                  help="Informações adicionais sobre o contrato")
        
        st.markdown("---")
        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            submitted_contrato = st.form_submit_button(
                "📄 Gerar Contrato (DOCX e PDF)",
                use_container_width=True,
                type="primary"
            )
        with col_btn2:
            submitted_adendo = st.form_submit_button(
                "📑 Gerar Adendo 2026 (DOCX e PDF)",
                use_container_width=True,
                type="secondary"
            )

    submitted = submitted_contrato or submitted_adendo
    
    # Processar formulário
    if submitted:
        if not all([tipo_servico, nome_servico, cnpj, valor, data_inicio, local_execucao, funcoes]):
            st.error("⚠️ Por favor, preencha todos os campos obrigatórios marcados com *")
        else:
            with st.container():
                st.markdown("---")
                
                with st.spinner("Gerando contrato..."):
                    modelo_selecionado = "Adendo 2026" if submitted_adendo else "Contrato de Servico"
                    docx_bytes, nome_docx, pdf_bytes, nome_pdf, motor_pdf = preencher_contrato(
                        tipo_servico=tipo_servico,
                        nome_servico=nome_servico,
                        cnpj=cnpj,
                        ie_cliente=ie_cliente if ie_cliente else "NÃO INFORMADO",
                        valor=valor,
                        data_inicio=data_inicio,
                        local_execucao=local_execucao,
                        funcoes=funcoes,
                        observacoes=observacoes,
                        modelo_path=MODELOS_DISPONIVEIS[modelo_selecionado],
                        modelo_nome=modelo_selecionado,
                        dados_cnpj=st.session_state.dados_cnpj
                    )
                
                if docx_bytes and nome_docx:
                    st.success("✅ Contrato gerado com sucesso!")
                    if motor_pdf == "Compativel (fallback)":
                        st.warning("⚠️ PDF gerado em modo compatível. O conteúdo foi preservado, mas a formatação pode ser simplificada.")
                    elif motor_pdf:
                        st.info(f"ℹ️ PDF gerado com: {motor_pdf}")
                    
                    col_down1, col_down2, col_down3 = st.columns([1, 1, 1])
                    
                    with col_down1:
                        st.download_button(
                            label="📥 Baixar DOCX",
                            data=docx_bytes,
                            file_name=nome_docx,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            type="secondary"
                        )
                    
                    with col_down2:
                        if pdf_bytes and nome_pdf:
                            st.download_button(
                                label="📄 Baixar PDF",
                                data=pdf_bytes,
                                file_name=nome_pdf,
                                mime="application/pdf",
                                use_container_width=True,
                                type="primary"
                            )
                        else:
                            st.info("PDF não disponível")
                    
                    with col_down3:
                        st.button("🔄 Novo Contrato", 
                                 use_container_width=True,
                                 on_click=lambda: st.session_state.clear())

if __name__ == "__main__":
    main()
